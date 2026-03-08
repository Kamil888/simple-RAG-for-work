import io
from docx import Document

CHUNK_SIZE = 2000
CHUNK_OVERLAP = 200
CHARS_PER_PAGE = 3000


def _split_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return [c for c in chunks if c.strip()]


def parse_docx(file_bytes: bytes, filename: str) -> list[dict]:
    doc = Document(io.BytesIO(file_bytes))

    blocks: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            blocks.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                blocks.append(row_text)

    full_text = "\n".join(blocks)
    chunks = []
    cumulative_chars = 0
    for chunk_index, fragment in enumerate(_split_text(full_text)):
        estimated_page = max(1, cumulative_chars // CHARS_PER_PAGE + 1)
        chunks.append({
            "text": fragment,
            "source": filename,
            "page": estimated_page,
            "chunk_index": chunk_index,
        })
        cumulative_chars += len(fragment)

    return chunks
